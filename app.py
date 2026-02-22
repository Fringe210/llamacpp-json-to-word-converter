import os
import re
import json
import tempfile
from datetime import datetime
from flask import Flask, render_template, request, send_file, flash, redirect, url_for
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

app = Flask(__name__)
app.secret_key = 'chiave_segreta_per_sessioni'

# Carica le traduzioni
with open('translations.json', 'r', encoding='utf-8') as f:
    TRANSLATIONS = json.load(f)

# Cartelle temporanee
UPLOAD_FOLDER = tempfile.gettempdir()
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

ALLOWED_EXTENSIONS = {'json'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def format_timestamp(ts):
    try:
        dt = datetime.fromtimestamp(ts / 1000)
        return dt.strftime("%d/%m/%Y %H:%M:%S")
    except:
        return str(ts)

def get_text(key, lang='it'):
    return TRANSLATIONS.get(lang, TRANSLATIONS['it']).get(key, key)

# Rimuove i blocchi di reasoning content generati dal modello
def strip_reasoning_content(text):
    if not text:
        return ""
    # Rimuove blocchi del tipo <<<reasoning_content_start>>>...<<<reasoning_content_end>>>
    text = re.sub(r'<<<reasoning_content_start>>>.*?<<<reasoning_content_end>>>', '', text, flags=re.DOTALL)
    # Rimuove eventuali tag rimasti senza chiusura (fino a fine stringa)
    text = re.sub(r'<<<reasoning_content_start>>>.*', '', text, flags=re.DOTALL)
    return text.strip()

# Funzione per pulire il testo (mantenuta per compatibilità, ora non rimuove più i marcatori markdown)
def clean_text(text):
    if not text:
        return ""
    text = strip_reasoning_content(text)
    return text

# Converte le espressioni LaTeX in simboli Unicode leggibili
LATEX_SYMBOLS = {
    r'\Psi': 'Ψ', r'\psi': 'ψ', r'\Phi': 'Φ', r'\phi': 'φ',
    r'\alpha': 'α', r'\beta': 'β', r'\gamma': 'γ', r'\delta': 'δ',
    r'\Delta': 'Δ', r'\Gamma': 'Γ', r'\Lambda': 'Λ', r'\lambda': 'λ',
    r'\mu': 'μ', r'\nu': 'ν', r'\pi': 'π', r'\Pi': 'Π',
    r'\rho': 'ρ', r'\sigma': 'σ', r'\Sigma': 'Σ', r'\tau': 'τ',
    r'\theta': 'θ', r'\Theta': 'Θ', r'\omega': 'ω', r'\Omega': 'Ω',
    r'\epsilon': 'ε', r'\eta': 'η', r'\xi': 'ξ', r'\zeta': 'ζ',
    r'\chi': 'χ', r'\kappa': 'κ', r'\iota': 'ι', r'\upsilon': 'υ',
    r'\infty': '∞', r'\partial': '∂', r'\nabla': '∇',
    r'\sum': '∑', r'\prod': '∏', r'\int': '∫',
    r'\sqrt': '√', r'\pm': '±', r'\mp': '∓',
    r'\times': '×', r'\div': '÷', r'\cdot': '·',
    r'\leq': '≤', r'\geq': '≥', r'\neq': '≠',
    r'\approx': '≈', r'\equiv': '≡', r'\sim': '∼',
    r'\in': '∈', r'\notin': '∉', r'\subset': '⊂', r'\supset': '⊃',
    r'\cup': '∪', r'\cap': '∩', r'\emptyset': '∅',
    r'\forall': '∀', r'\exists': '∃', r'\neg': '¬',
    r'\wedge': '∧', r'\vee': '∨', r'\oplus': '⊕',
    r'\rightarrow': '→', r'\leftarrow': '←', r'\Rightarrow': '⇒',
    r'\Leftarrow': '⇐', r'\leftrightarrow': '↔', r'\Leftrightarrow': '⇔',
    r'\uparrow': '↑', r'\downarrow': '↓',
    r'\hbar': 'ℏ', r'\ell': 'ℓ', r'\Re': 'ℜ', r'\Im': 'ℑ',
    r'\circ': '∘', r'\bullet': '•', r'\ldots': '…', r'\cdots': '⋯',
    r'\^': '^', r'\_': '_',
}

def latex_to_unicode(text):
    """Converte espressioni LaTeX $...$ in testo Unicode leggibile."""
    def convert_expr(expr):
        # Sostituisce comandi LaTeX noti con simboli Unicode
        for cmd, sym in sorted(LATEX_SYMBOLS.items(), key=lambda x: -len(x[0])):
            expr = expr.replace(cmd, sym)
        # Gestisce ^{...} → apici (usa caratteri superscript dove possibile)
        sup_map = {'0':'⁰','1':'¹','2':'²','3':'³','4':'⁴','5':'⁵','6':'⁶','7':'⁷','8':'⁸','9':'⁹','+':'⁺','-':'⁻','n':'ⁿ'}
        def sup_replace(m):
            content = m.group(1)
            return ''.join(sup_map.get(c, c) for c in content)
        expr = re.sub(r'\^\{([^}]+)\}', sup_replace, expr)
        expr = re.sub(r'\^(\w)', lambda m: ''.join({'0':'⁰','1':'¹','2':'²','3':'³','4':'⁴','5':'⁵','6':'⁶','7':'⁷','8':'⁸','9':'⁹'}.get(c,c) for c in m.group(1)), expr)
        # Gestisce _{...} → pedici
        sub_map = {'0':'₀','1':'₁','2':'₂','3':'₃','4':'₄','5':'₅','6':'₆','7':'₇','8':'₈','9':'₉','+':'₊','-':'₋','n':'ₙ'}
        def sub_replace(m):
            content = m.group(1)
            return ''.join(sub_map.get(c, c) for c in content)
        expr = re.sub(r'_\{([^}]+)\}', sub_replace, expr)
        expr = re.sub(r'_(\w)', lambda m: ''.join(sub_map.get(c,c) for c in m.group(1)), expr)
        # Gestisce \frac{a}{b} → a/b
        expr = re.sub(r'\\frac\{([^}]+)\}\{([^}]+)\}', r'\1/\2', expr)
        # Rimuove \left \right e altri comandi di dimensionamento
        expr = re.sub(r'\\(left|right|big|Big|bigg|Bigg)', '', expr)
        # Rimuove eventuali {} rimasti
        expr = expr.replace('{', '').replace('}', '')
        # Rimuove backslash rimasti seguiti da spazio o fine
        expr = re.sub(r'\\(?=\s|$)', '', expr)
        return expr.strip()

    # Sostituisce $...$  (inline math) con il testo convertito
    result = re.sub(r'\$\$(.+?)\$\$', lambda m: convert_expr(m.group(1)), text, flags=re.DOTALL)
    result = re.sub(r'\$(.+?)\$', lambda m: convert_expr(m.group(1)), result)
    return result

# Parsa il markdown inline (grassetto e corsivo) e aggiunge run formattati al paragrafo
def add_inline_markdown(paragraph, text, base_size=11, base_bold=False, base_italic=False, color=None):
    # Prima converti il LaTeX in Unicode
    text = latex_to_unicode(text)
    # Pattern per **bold**, *italic*, ***bold+italic***
    pattern = re.compile(r'(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*)')
    last_end = 0
    for m in pattern.finditer(text):
        # Testo prima del match
        if m.start() > last_end:
            run = paragraph.add_run(text[last_end:m.start()])
            run.font.size = Pt(base_size)
            run.font.bold = base_bold
            run.font.italic = base_italic
            if color:
                run.font.color.rgb = color
        # Testo formattato
        if m.group(2):  # ***bold+italic***
            run = paragraph.add_run(m.group(2))
            run.font.bold = True
            run.font.italic = True
        elif m.group(3):  # **bold**
            run = paragraph.add_run(m.group(3))
            run.font.bold = True
            run.font.italic = base_italic
        elif m.group(4):  # *italic*
            run = paragraph.add_run(m.group(4))
            run.font.bold = base_bold
            run.font.italic = True
        run.font.size = Pt(base_size)
        if color:
            run.font.color.rgb = color
        last_end = m.end()
    # Testo rimanente
    if last_end < len(text):
        run = paragraph.add_run(text[last_end:])
        run.font.size = Pt(base_size)
        run.font.bold = base_bold
        run.font.italic = base_italic
        if color:
            run.font.color.rgb = color

# Aggiunge una riga di testo al doc gestendo heading markdown e inline markdown
def add_markdown_line(doc, line, indent=None, italic=False):
    stripped = line.strip()
    if not stripped:
        return

    # Heading 1: # testo
    if re.match(r'^#{1}\s+', stripped) and not re.match(r'^#{2,}', stripped):
        heading_text = re.sub(r'^#+\s+', '', stripped)
        heading_text = latex_to_unicode(heading_text)
        p = doc.add_heading(heading_text, level=1)
        return

    # Heading 2: ## testo
    if re.match(r'^#{2}\s+', stripped) and not re.match(r'^#{3,}', stripped):
        heading_text = re.sub(r'^#+\s+', '', stripped)
        heading_text = latex_to_unicode(heading_text)
        p = doc.add_heading(heading_text, level=2)
        return

    # Heading 3+: ### testo
    if re.match(r'^#{3,}\s+', stripped):
        heading_text = re.sub(r'^#+\s+', '', stripped)
        heading_text = latex_to_unicode(heading_text)
        p = doc.add_heading(heading_text, level=3)
        return

    # Punto elenco: "* testo" o "- testo" o "• testo"
    bullet_match = re.match(r'^(\*|-|•)\s+(.+)$', stripped)
    if bullet_match:
        bullet_text = bullet_match.group(2)
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        if indent:
            p.paragraph_format.left_indent = indent
        add_inline_markdown(p, bullet_text, base_size=11, base_italic=italic)
        return p

    # Punto elenco numerato: "1. testo" o "1) testo"
    numbered_match = re.match(r'^\d+[.)]\s+(.+)$', stripped)
    if numbered_match:
        item_text = numbered_match.group(1)
        p = doc.add_paragraph(style='List Number')
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        if indent:
            p.paragraph_format.left_indent = indent
        add_inline_markdown(p, item_text, base_size=11, base_italic=italic)
        return p

    # Paragrafo normale con inline markdown
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    if indent:
        p.paragraph_format.left_indent = indent
    add_inline_markdown(p, stripped, base_size=11, base_italic=italic)
    return p

# Funzione per rilevare e parsare tabelle Markdown
def is_markdown_table(line):
    return line.strip().startswith('|') and line.strip().endswith('|')

def parse_markdown_table(content):
    """
    Rileva se il contenuto contiene tabelle Markdown e le парсит.
    Restituisce una lista di dizionari: {'type': 'table', 'data': [...]}
    oppure {'type': 'text', 'data': ...}
    """
    if not content:
        return [{'type': 'text', 'data': content}]
    
    lines = content.split('\n')
    result = []
    current_text_lines = []
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        # Rileva inizio tabella Markdown
        if is_markdown_table(line):
            # Salva il testo accumulato prima della tabella
            if current_text_lines:
                text_block = '\n'.join(current_text_lines).strip()
                if text_block:
                    result.append({'type': 'text', 'data': text_block})
                current_text_lines = []
            
            # Parsa la tabella
            table_data = []
            # Leggi le righe della tabella fino a quando ci sono righe che iniziano con |
            while i < len(lines) and is_markdown_table(lines[i]):
                row = [cell.strip() for cell in lines[i].strip().split('|')[1:-1]]  # Rimuovi primo e ultimo elemento vuoto
                
                # Salta la riga di separazione (---|---|)
                if all(cell.strip() in ('', '-', ':', '---', ':--', '--:', ':-:', ':---', '---:', ':---:') for cell in row):
                    i += 1
                    continue
                
                if row:
                    table_data.append(row)
                i += 1
            
            if table_data:
                result.append({'type': 'table', 'data': table_data})
            
            # Continua il ciclo senza incrementare i perché già fatto nel while
            continue
        else:
            if line:  # Solo linee non vuote
                current_text_lines.append(lines[i])
            elif current_text_lines:
                # Linea vuota: salva il testo accumulato
                text_block = '\n'.join(current_text_lines).strip()
                if text_block:
                    result.append({'type': 'text', 'data': text_block})
                current_text_lines = []
        
        i += 1
    
    # Salva eventuale testo rimanente
    if current_text_lines:
        text_block = '\n'.join(current_text_lines).strip()
        if text_block:
            result.append({'type': 'text', 'data': text_block})
    
    # Se non è stata trovata nessuna tabella, restituisci tutto come testo
    if not any(block['type'] == 'table' for block in result):
        return [{'type': 'text', 'data': content}]
    
    return result

# Funzione per aggiungere una tabella al documento Word
def add_table_to_doc(doc, table_data, style='Table Grid'):
    if not table_data:
        return
    
    # La prima riga è l'intestazione
    headers = table_data[0]
    rows = table_data[1:] if len(table_data) > 1 else []
    
    # Crea la tabella
    num_rows = len(rows) + 1
    num_cols = len(headers)
    
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = style
    
    # Aggiungi l'intestazione
    for col_idx, header in enumerate(headers):
        cell = table.rows[0].cells[col_idx]
        cell.text = header
        # Formatta l'intestazione
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(11)
    
    # Aggiungi le righe dati
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_data in enumerate(row_data):
            if col_idx < num_cols:
                cell = table.rows[row_idx + 1].cells[col_idx]
                cell.text = cell_data
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)

def convert_json_to_docx(json_data, doc, options, lang):
    t = lambda k: get_text(k, lang)
    
    # Informazioni sulla conversazione
    conv = json_data.get('conv', {})
    
    # Titolo del documento
    title = doc.add_heading(t('doc_title'), 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Informazioni della conversazione
    doc.add_paragraph()
    p_info = doc.add_paragraph()
    run = p_info.add_run(t('doc_info'))
    run.bold = True
    run.font.size = Pt(14)
    
    # Tabella metadati
    table_meta = doc.add_table(rows=4, cols=2)
    table_meta.style = 'Table Grid'
    
    meta_data = [
        (t('conv_id'), conv.get('id', 'N/A')),
        (t('conv_name'), conv.get('name', 'N/A')),
        (t('conv_last_mod'), format_timestamp(conv.get('lastModified', 0))),
        (t('conv_node'), conv.get('currNode', 'N/A')[:20] + '...' if len(conv.get('currNode', '')) > 20 else conv.get('currNode', 'N/A'))
    ]
    
    for i, (label, value) in enumerate(meta_data):
        row = table_meta.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
        for cell in row.cells:
            if cell.paragraphs and cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].font.bold = True

    if options.get('show_divider'):
        doc.add_paragraph()
        doc.add_paragraph('━' * 60)
        doc.add_paragraph()

    # Processa i messaggi
    messages = json_data.get('messages', [])
    
    # Determina nomi personalizzati o default
    user_label = options.get('custom_user_name', '').strip() or t('default_user')
    assistant_label = options.get('custom_assistant_name', '').strip() or t('default_assistant')

    for idx, msg in enumerate(messages, 1):
        role = msg.get('role', 'unknown')
        content = msg.get('content', '')
        msg_type = msg.get('type', 'text')
        timestamp = msg.get('timestamp', 0)
        
        if not content and msg_type != 'text':
            continue
        
        # Colori
        if role == 'user':
            role_display = user_label
            role_color = RGBColor(33, 150, 243)
        elif role == 'assistant':
            role_display = assistant_label
            role_color = RGBColor(76, 175, 80)
        else:
            role_display = f"📋 {role.upper()}"
            role_color = RGBColor(158, 158, 158)
        
        # Intestazione del messaggio
        doc.add_paragraph()
        p = doc.add_paragraph()
        
        # Numero messaggio
        if options.get('show_numbers'):
            run_num = p.add_run(f"[{idx}] ")
            run_num.font.size = Pt(10)
            run_num.font.color.rgb = RGBColor(158, 158, 158)
        
        run_role = p.add_run(f"{role_display}")
        run_role.font.size = Pt(12)
        run_role.font.bold = True
        run_role.font.color.rgb = role_color
        
        # Data e Ora
        if options.get('show_date'):
            run_time = p.add_run(f" • {format_timestamp(timestamp)}")
            run_time.font.size = Pt(9)
            run_time.font.color.rgb = RGBColor(158, 158, 158)
        
        # Modello AI
        if options.get('show_model') and role == 'assistant' and msg.get('model'):
            p_model = doc.add_paragraph()
            run_model = p_model.add_run(f"   📡 {msg.get('model')}")
            run_model.font.size = Pt(9)
            run_model.font.italic = True
            run_model.font.color.rgb = RGBColor(158, 158, 158)
        
        # Contenuto del messaggio
        if content:
            content_clean = content.replace('\\n', '\n').replace('\\t', '\t')
            content_clean = clean_text(content_clean)
            
            # Parse del contenuto per rilevare tabelle
            content_blocks = parse_markdown_table(content_clean)
            
            for block in content_blocks:
                if block['type'] == 'table':
                    # Aggiungi la tabella
                    add_table_to_doc(doc, block['data'])
                else:
                    # Testo normale con supporto markdown
                    for line in block['data'].split('\n'):
                        if line.strip():
                            add_markdown_line(doc, line)
        
        # Contenuto Extra
        extra = msg.get('extra', [])
        if extra:
            doc.add_paragraph()
            p_extra = doc.add_paragraph()
            run_extra = p_extra.add_run(t('extra_content'))
            run_extra.font.bold = True
            run_extra.font.size = Pt(10)
            
            for item in extra:
                if isinstance(item, dict):
                    if item.get('type') == 'TEXT' and item.get('name') == 'Pasted':
                        extra_content = item.get('content', '')
                        if extra_content:
                            extra_clean = extra_content.replace('\\n', '\n')
                            extra_clean = clean_text(extra_clean)
                            
                            # Anche per l'extra content, rileva tabelle
                            extra_blocks = parse_markdown_table(extra_clean)
                            
                            for block in extra_blocks:
                                if block['type'] == 'table':
                                    add_table_to_doc(doc, block['data'])
                                else:
                                    for line in block['data'].split('\n'):
                                        if line.strip():
                                            add_markdown_line(doc, line, indent=Inches(0.3), italic=True)
        
        # Dati Prompt (Timing)
        if options.get('show_prompt'):
            timings = msg.get('timings', {})
            if timings and role == 'assistant':
                p_timing = doc.add_paragraph()
                run_timing = p_timing.add_run("⏱️ ")
                run_timing.font.size = Pt(9)
                
                timing_text = f"Prompt: {timings.get('prompt_n', 'N/A')} token ({timings.get('prompt_ms', 0):.1f}ms) | "
                timing_text += f"Output: {timings.get('predicted_n', 'N/A')} token ({timings.get('predicted_ms', 0):.1f}ms)"
                
                run_timing_val = p_timing.add_run(timing_text)
                run_timing_val.font.size = Pt(9)
                run_timing_val.font.color.rgb = RGBColor(158, 158, 158)
        
        # Separatore finale
        if options.get('show_divider'):
            doc.add_paragraph()
            p_sep = doc.add_paragraph('─' * 60)
            p_sep.alignment = WD_ALIGN_PARAGRAPH.CENTER

def process_json(json_file_path, options, lang):
    with open(json_file_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    convert_json_to_docx(data, doc, options, lang)
    
    doc.add_paragraph()
    if options.get('show_divider'):
        doc.add_paragraph('━' * 60)
    p_footnote = doc.add_paragraph()
    run_footnote = p_footnote.add_run("📄 " + get_text('generated_at', lang))
    run_footnote.font.size = Pt(9)
    run_footnote.font.italic = True
    run_footnote.font.color.rgb = RGBColor(158, 158, 158)
    
    return doc

@app.route('/')
def index():
    # Passa le lingue disponibili al template
    return render_template('index.html', languages=TRANSLATIONS.keys())

@app.route('/convert', methods=['POST'])
def convert():
    try:
        if 'file' not in request.files:
            flash('❌ Nessun file caricato', 'error')
            return redirect(url_for('index'))
        
        file = request.files['file']
        
        if file.filename == '':
            flash('❌ Nessun file selezionato', 'error')
            return redirect(url_for('index'))
        
        if file and allowed_file(file.filename):
            # Estrai opzioni
            options = {
                'show_date': request.form.get('show_date') == 'on',
                'show_divider': request.form.get('show_divider') == 'on',
                'show_model': request.form.get('show_model') == 'on',
                'show_prompt': request.form.get('show_prompt') == 'on',
                'show_numbers': request.form.get('show_numbers') == 'on',
                'custom_user_name': request.form.get('custom_user_name', ''),
                'custom_assistant_name': request.form.get('custom_assistant_name', '')
            }
            lang = request.form.get('language', 'it')
            
            filename = f"input_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(filepath)
            
            try:
                doc = process_json(filepath, options, lang)
                
                output_filename = f"conversazione_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
                output_path = os.path.join(app.config['UPLOAD_FOLDER'], output_filename)
                
                doc.save(output_path)
                
                flash('✅ Conversione completata con successo!', 'success')
                
                return send_file(
                    output_path,
                    as_attachment=True,
                    download_name=output_filename,
                    mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                )
                
            except Exception as e:
                flash(f'❌ Errore durante l\'elaborazione: {str(e)}', 'error')
                return redirect(url_for('index'))
            
            finally:
                if os.path.exists(filepath):
                    os.remove(filepath)
        else:
            flash('❌ Formato file non consentito.', 'error')
            return redirect(url_for('index'))
    
    except Exception as e:
        flash(f'❌ Errore: {str(e)}', 'error')
        return redirect(url_for('index'))

@app.route('/sample')
def download_sample():
    sample_json = {
        "conv": {
            "id": "sample-id",
            "name": "esempio",
            "lastModified": 1771702156904,
            "currNode": "sample-node"
        },
        "messages": [
            {
                "convId": "sample-id",
                "role": "user",
                "content": "Ciao! Questo è un messaggio di esempio.",
                "type": "text",
                "timestamp": 1771702156956
            },
            {
                "convId": "sample-id",
                "role": "assistant",
                "content": "Ecco una tabella di esempio:\n\n| Nome | Età | Città |\n|------|-----|-------|\n| Mario | 30 | Roma |\n| Giulia | 25 | Milano |\n| Luca | 35 | Napoli |",
                "type": "text",
                "timestamp": 1771702156981,
                "model": "Modello-Esempio",
                "timings": {"prompt_n": 10, "prompt_ms": 50.5, "predicted_n": 20, "predicted_ms": 100.2}
            }
        ]
    }
    
    sample_path = os.path.join(app.config['UPLOAD_FOLDER'], 'sample_conversation.json')
    with open(sample_path, 'w', encoding='utf-8') as f:
        json.dump(sample_json, f, indent=2, ensure_ascii=False)
    
    return send_file(
        sample_path,
        as_attachment=True,
        download_name='sample_conversation.json',
        mimetype='application/json'
    )

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)
